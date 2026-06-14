from __future__ import annotations

from pathlib import Path


LAYOUT_DIRS = (
    "user",
    "user/loop-memory",
    "data/raw",
    "data/processed",
    "references",
    "knowledge_base",
    "knowledge_base/tables",
    "knowledge_base/indexes",
    "artifacts",
    "plots",
    "tools",
    "tools/generated",
    "runs/iterations",
    "reports",
    "reports/iterations",
    "logs/nodes",
    "state/nodes",
    "training",
)


def ensure_workspace_layout(root: Path) -> None:
    for rel in LAYOUT_DIRS:
        (root / rel).mkdir(parents=True, exist_ok=True)
    ensure_builtin_tools(root)
    ensure_reference_text_derivatives(root)


def ensure_reference_text_derivatives(root: Path) -> None:
    references = root / "references"
    if not references.exists():
        return
    for path in references.iterdir():
        if path.is_file() and path.suffix.lower() == ".docx":
            write_reference_text_derivative(root, path)


def write_reference_text_derivative(root: Path, path: Path) -> str | None:
    if path.suffix.lower() != ".docx":
        return None
    target = path.with_suffix(path.suffix + ".txt")
    if target.exists() and target.stat().st_mtime >= path.stat().st_mtime:
        return str(target.relative_to(root))
    try:
        text = extract_docx_text(path)
    except Exception:
        return None
    if not text.strip():
        return None
    target.write_text(text, encoding="utf-8")
    return str(target.relative_to(root))


def ensure_builtin_tools(root: Path) -> None:
    script = root / "tools" / "read_docx.py"
    if script.exists():
        return
    script.write_text(READ_DOCX_SCRIPT, encoding="utf-8")


def extract_docx_text(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        chunks.append(f"\n[Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    return "\n\n".join(chunks).strip() + "\n"


READ_DOCX_SCRIPT = '''from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def extract_docx(path: Path) -> str:
    document = Document(str(path))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        chunks.append(f"\\n[Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\\n", " ") for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    return "\\n\\n".join(chunks)


def main() -> None:
    if len(sys.argv) not in {2, 3}:
        raise SystemExit("Usage: uv run python tools/read_docx.py <path-to-docx> [output.txt]")
    path = Path(sys.argv[1])
    text = extract_docx(path)
    if len(sys.argv) == 3:
        out = Path(sys.argv[2])
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(text + "\\n", encoding="utf-8")
        print(str(out))
    else:
        print(text)


if __name__ == "__main__":
    main()
'''
