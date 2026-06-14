from __future__ import annotations

import zipfile
from datetime import UTC, datetime
from pathlib import Path
import shutil
from typing import Any
from uuid import uuid4

from harnessing_ts.schema import CONTROL_MODES, NODE_TYPES, ControlRequest, NodeSession, NodeType, Part, RunRecord, RuntimeSettings, TimelineEvent, WorkspaceState
from harnessing_ts.state.jsonl import append_jsonl, clear_file, read_json, read_jsonl, write_json
from harnessing_ts.workspace_runtime import ensure_workspace_uv_environment


def now_iso() -> str:
    return datetime.now(UTC).isoformat().replace("+00:00", "Z")


class WorkspaceStore:
    def __init__(self, root: Path) -> None:
        self.root = root
        self.state_path = root / "state" / "workspace.json"
        self.timeline_path = root / "logs" / "timeline.jsonl"
        self.main_log_path = root / "logs" / "main.jsonl"
        self.knowledge_graph_log_path = root / "logs" / "knowledge-graph-builder.jsonl"
        self.knowledge_reasoning_log_path = root / "logs" / "knowledge-reasoning.jsonl"
        self.runtime_path = root / "state" / "runtime.json"
        self.runtime_settings_path = root / "state" / "runtime-settings.json"
        self.knowledge_graph_llm_path = root / "state" / "knowledge-graph-llm.json"
        self.knowledge_graph_status_path = root / "state" / "knowledge-graph-build.json"
        self.knowledge_graph_path = root / "artifacts" / "knowledge-graph.json"
        self.node_log_dir = root / "logs" / "nodes"
        self.node_meta_dir = root / "state" / "nodes"

    def initialize(self, mode: str = "manual") -> WorkspaceState:
        self.ensure_layout()
        control_mode = "auto" if mode == "auto" else "manual"
        runtime_status = ensure_workspace_uv_environment(self.root)
        if runtime_status.get("state") == "failed":
            self.append_timeline({
                "type": "workspace_runtime_failed",
                "timestamp": now_iso(),
                "message": runtime_status.get("message", "workspace uv environment failed"),
                "payload": runtime_status,
            })
        existing = self.read_state()
        if existing:
            if existing.get("mode") not in CONTROL_MODES:
                existing["mode"] = control_mode
                self.write_state(existing)
            if existing.get("controlMode") != control_mode:
                existing["controlMode"] = control_mode
                existing["mode"] = control_mode
                self.write_state(existing)
            if "pendingControl" not in existing:
                existing["pendingControl"] = None
                self.write_state(existing)
            if "activeNodeSessionId" not in existing:
                existing["activeNodeSessionId"] = None
                self.write_state(existing)
            if "contractConfirmed" not in existing:
                existing["contractConfirmed"] = bool(existing.pop("taskContractConfirmed", False))
                self.write_state(existing)
            if "finalSummaryConfirmed" not in existing:
                existing["finalSummaryConfirmed"] = bool(existing.pop("finalSolutionConfirmed", False))
                self.write_state(existing)
            settings = self.read_runtime_settings()
            if existing.get("runtimeSettings") != settings:
                existing["runtimeSettings"] = settings
                self.write_state(existing)
            if "pendingHumanGate" in existing:
                existing.pop("pendingHumanGate", None)
                self.write_state(existing)
            valid_nodes = set(NODE_TYPES)
            filtered_completed = [node for node in existing.get("completedNodes", []) if node in valid_nodes]
            if filtered_completed != existing.get("completedNodes", []):
                existing["completedNodes"] = filtered_completed
                self.write_state(existing)
            if existing.get("activeNode") and existing["activeNode"] not in valid_nodes:
                existing["activeNode"] = None
                existing["activeNodeSessionId"] = None
                self.write_state(existing)
            return existing
        ts = now_iso()
        state: WorkspaceState = {
            "workspaceId": str(uuid4()),
            "workspacePath": str(self.root),
            "createdAt": ts,
            "updatedAt": ts,
            "mode": control_mode,
            "controlMode": control_mode,
            "pendingControl": None,
            "activeNode": None,
            "activeNodeSessionId": None,
            "completedNodes": [],
            "contractConfirmed": False,
            "finalSummaryConfirmed": False,
            "runtimeSettings": self.read_runtime_settings(),
        }
        self.write_state(state)
        self.append_timeline({"type": "workspace_initialized", "timestamp": ts, "payload": {"mode": state["mode"]}})
        return state

    def set_pending_control(self, request: ControlRequest) -> ControlRequest:
        state = self.read_state()
        if not state:
            raise RuntimeError("Workspace state is not initialized.")
        state["pendingControl"] = request
        self.write_state(state)
        self.append_timeline({
            "type": "control_pending",
            "timestamp": now_iso(),
            "nodeSessionId": request.get("nodeSessionId"),
            "nodeType": request.get("nodeType"),
            "message": request.get("message") or request["kind"],
            "payload": request,
        })
        return request

    def clear_pending_control(self) -> ControlRequest | None:
        state = self.read_state()
        if not state:
            return None
        request = state.get("pendingControl")
        state["pendingControl"] = None
        self.write_state(state)
        return request

    def ensure_layout(self) -> None:
        for rel in (
            "user",
            "user/loop-memory",
            "data/raw",
            "data/processed",
            "references",
            "knowledge_base",
            "knowledge_base/tables",
            "knowledge_base/indexes",
            "artifacts",
            "plots",
            "tools",
            "tools/generated",
            "runs/iterations",
            "reports",
            "reports/iterations",
            "logs/nodes",
            "state/nodes",
            "training",
        ):
            (self.root / rel).mkdir(parents=True, exist_ok=True)
        self.ensure_builtin_tools()
        self.ensure_reference_text_derivatives()

    def read_state(self) -> WorkspaceState | None:
        return read_json(self.state_path)

    def read_runtime_status(self) -> dict[str, Any] | None:
        return read_json(self.runtime_path)

    def read_knowledge_graph_build_status(self) -> dict[str, Any]:
        return read_json(self.knowledge_graph_status_path) or {
            "running": False,
            "status": "idle",
            "startedAt": None,
            "finishedAt": None,
            "trigger": None,
            "message": "Knowledge graph builder has not run yet.",
        }

    def write_knowledge_graph_build_status(self, status: dict[str, Any]) -> dict[str, Any]:
        current = self.read_knowledge_graph_build_status()
        current.update(status)
        write_json(self.knowledge_graph_status_path, current)
        return current

    @property
    def main_llm_path(self) -> Path:
        return self.root / "config.llm.json"

    def read_main_llm_config(self) -> dict[str, Any]:
        raw = read_json(self.main_llm_path) or {}
        return _sanitize_llm_config(raw, default_auth_mode="manual")

    def write_main_llm_config(self, values: dict[str, Any]) -> dict[str, Any]:
        current = self.read_main_llm_config()
        merged = dict(current)
        for key in ("authMode", "protocol", "model", "apiKey", "baseUrl", "contextWindow"):
            if key not in values:
                continue
            value = values[key]
            if key == "apiKey" and (value is None or value == "" or _looks_masked_secret(str(value))):
                continue
            merged[key] = value
        sanitized = _sanitize_llm_config(merged, default_auth_mode="manual")
        self.main_llm_path.parent.mkdir(parents=True, exist_ok=True)
        write_json(self.main_llm_path, sanitized)
        self.append_timeline({
            "type": "main_llm_updated",
            "timestamp": now_iso(),
            "message": f"model={sanitized.get('model') or 'sdk-default'}",
            "payload": _mask_llm_config_dict(sanitized),
        })
        return sanitized

    def read_knowledge_graph_llm_config(self) -> dict[str, Any]:
        raw = read_json(self.knowledge_graph_llm_path) or {}
        return _sanitize_llm_config(raw, default_auth_mode="manual")

    def write_knowledge_graph_llm_config(self, values: dict[str, Any]) -> dict[str, Any]:
        current = self.read_knowledge_graph_llm_config()
        merged = dict(current)
        for key in ("authMode", "protocol", "model", "apiKey", "baseUrl", "contextWindow"):
            if key not in values:
                continue
            value = values[key]
            if key == "apiKey" and (value is None or value == "" or _looks_masked_secret(str(value))):
                continue
            merged[key] = value
        sanitized = _sanitize_llm_config(merged, default_auth_mode="manual")
        write_json(self.knowledge_graph_llm_path, sanitized)
        self.append_timeline({
            "type": "knowledge_graph_llm_updated",
            "timestamp": now_iso(),
            "message": f"model={sanitized.get('model') or 'sdk-default'}",
            "payload": _mask_llm_config_dict(sanitized),
        })
        return sanitized

    def read_runtime_settings(self) -> RuntimeSettings:
        raw = read_json(self.runtime_settings_path) or {}
        return {
            "iterativeCandidateCount": _bounded_int(raw.get("iterativeCandidateCount"), default=3, minimum=1, maximum=8),
            "knowledgeGraphExtractionDepth": _bounded_int(raw.get("knowledgeGraphExtractionDepth"), default=2, minimum=1, maximum=4),
        }

    def write_runtime_settings(self, settings: dict[str, Any]) -> RuntimeSettings:
        current = self.read_runtime_settings()
        if "iterativeCandidateCount" in settings:
            current["iterativeCandidateCount"] = _bounded_int(settings.get("iterativeCandidateCount"), default=current["iterativeCandidateCount"], minimum=1, maximum=8)
        if "knowledgeGraphExtractionDepth" in settings:
            current["knowledgeGraphExtractionDepth"] = _bounded_int(settings.get("knowledgeGraphExtractionDepth"), default=current["knowledgeGraphExtractionDepth"], minimum=1, maximum=4)
        write_json(self.runtime_settings_path, current)
        state = self.read_state()
        if state:
            state["runtimeSettings"] = current
            self.write_state(state)
        self.append_timeline({
            "type": "runtime_settings_updated",
            "timestamp": now_iso(),
            "message": f"iterativeCandidateCount={current['iterativeCandidateCount']}, knowledgeGraphExtractionDepth={current['knowledgeGraphExtractionDepth']}",
            "payload": current,
        })
        return current

    def read_knowledge_graph(self) -> dict[str, Any]:
        from harnessing_ts.knowledge_graph import read_graph_view

        return read_graph_view(self.root)

    def read_knowledge_base_summary(self) -> dict[str, Any]:
        from harnessing_ts.knowledge_graph import read_knowledge_base_summary

        return read_knowledge_base_summary(self.root)

    def read_knowledge_base_cards(self, kind: str, limit: int = 200) -> dict[str, Any]:
        from harnessing_ts.knowledge_graph import read_knowledge_base_cards

        return read_knowledge_base_cards(self.root, kind, limit)

    def read_knowledge_graph_parts(self) -> list[Part]:
        return read_jsonl(self.knowledge_graph_log_path)

    def write_state(self, state: WorkspaceState) -> None:
        state["updatedAt"] = now_iso()
        write_json(self.state_path, state)

    def create_node_session(self, node_type: NodeType, rationale: str | None = None, input_summary: str | None = None) -> NodeSession:
        node: NodeSession = {
            "id": str(uuid4()),
            "nodeType": node_type,
            "status": "created",
            "startedAt": now_iso(),
        }
        if rationale:
            node["rationale"] = rationale
        if input_summary:
            node["inputSummary"] = input_summary
        self.write_node_session(node)
        return node

    def read_node_session(self, node_id: str) -> NodeSession | None:
        return read_json(self.node_meta_dir / f"{node_id}.json")

    def list_node_sessions(self) -> list[NodeSession]:
        try:
            names = sorted(self.node_meta_dir.iterdir())
        except FileNotFoundError:
            return []
        sessions: list[NodeSession] = []
        for path in names:
            if path.suffix != ".json":
                continue
            session = read_json(path)
            if session:
                sessions.append(session)
        return sorted(sessions, key=lambda item: item.get("startedAt", ""))

    def write_node_session(self, node: NodeSession) -> None:
        write_json(self.node_meta_dir / f"{node['id']}.json", node)

    def node_log_path(self, node_session_id: str) -> Path:
        return self.node_log_dir / f"{node_session_id}.jsonl"

    def append_main_part(self, part: Part) -> None:
        append_jsonl(self.main_log_path, part)

    def append_node_part(self, node_session_id: str, part: Part) -> None:
        append_jsonl(self.node_log_path(node_session_id), part)

    def append_timeline(self, event: TimelineEvent) -> None:
        append_jsonl(self.timeline_path, event)

    def record_artifact(self, path: str, node_session_id: str | None = None, node_type: NodeType | None = None, summary: str | None = None) -> None:
        self.append_timeline({
            "type": "artifact_recorded",
            "timestamp": now_iso(),
            "nodeSessionId": node_session_id,
            "nodeType": node_type,
            "message": summary or path,
            "payload": {"path": path},
        })

    def record_run(self, record: RunRecord) -> None:
        append_jsonl(self.root / "runs" / "registry.jsonl", record)
        self.append_timeline({
            "type": "run_recorded",
            "timestamp": now_iso(),
            "nodeSessionId": record.get("nodeSessionId"),
            "nodeType": record.get("nodeType"),
            "message": record.get("summary") or record.get("runId"),
            "payload": record,
        })

    def read_timeline(self) -> list[TimelineEvent]:
        return read_jsonl(self.timeline_path)

    def read_main_parts(self) -> list[Part]:
        return read_jsonl(self.main_log_path)

    def read_node_parts(self, node_session_id: str) -> list[Part]:
        return read_jsonl(self.node_log_path(node_session_id))

    def list_file_tree(self, max_entries: int = 800) -> dict[str, Any]:
        self.ensure_layout()
        count = 0

        def build(path: Path, rel: str) -> dict[str, Any] | None:
            nonlocal count
            if count >= max_entries:
                return None
            name = path.name if rel else self.root.name
            if path.name in {".git", ".venv", "__pycache__", "node_modules"}:
                return None
            count += 1
            stat = path.stat()
            if path.is_dir():
                children: list[dict[str, Any]] = []
                try:
                    entries = sorted(path.iterdir(), key=lambda item: (not item.is_dir(), item.name.lower()))
                except PermissionError:
                    entries = []
                for child in entries:
                    child_rel = str(child.relative_to(self.root))
                    child_node = build(child, child_rel)
                    if child_node:
                        children.append(child_node)
                return {
                    "name": name,
                    "path": rel,
                    "kind": "dir",
                    "mtime": datetime.fromtimestamp(stat.st_mtime, UTC).isoformat().replace("+00:00", "Z"),
                    "children": children,
                }
            return {
                "name": name,
                "path": rel,
                "kind": "file",
                "size": stat.st_size,
                "mtime": datetime.fromtimestamp(stat.st_mtime, UTC).isoformat().replace("+00:00", "Z"),
            }

        return {
            "root": str(self.root),
            "truncated": count >= max_entries,
            "tree": build(self.root, ""),
        }

    def read_text_file(self, rel_path: str, max_bytes: int = 256_000) -> dict[str, Any]:
        target = (self.root / rel_path).resolve()
        root = self.root.resolve()
        if target != root and root not in target.parents:
            raise ValueError("Path is outside workspace.")
        if not target.exists() or not target.is_file():
            raise FileNotFoundError(rel_path)
        if target.stat().st_size > max_bytes:
            return {"path": rel_path, "truncated": True, "text": "", "size": target.stat().st_size}
        try:
            text = target.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return {"path": rel_path, "binary": True, "text": "", "size": target.stat().st_size}
        return {"path": rel_path, "text": text, "size": target.stat().st_size}

    def write_reference_file(self, filename: str, content: bytes) -> str:
        safe_name = _safe_filename(filename)
        target = self.root / "references" / safe_name
        if target.exists():
            stem = target.stem
            suffix = target.suffix
            target = self.root / "references" / f"{stem}-{uuid4().hex[:8]}{suffix}"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_bytes(content)
        rel = str(target.relative_to(self.root))
        derivative = self._write_reference_text_derivative(target)
        self.append_timeline({
            "type": "reference_uploaded",
            "timestamp": now_iso(),
            "message": rel,
            "payload": {"path": rel, "size": len(content), "textDerivative": derivative},
        })
        return rel

    def extract_raw_data_zip(self, filename: str, content: bytes) -> dict[str, Any]:
        safe_name = _safe_filename(filename)
        if Path(safe_name).suffix.lower() != ".zip":
            raise ValueError("Only .zip archives can be uploaded as raw data.")

        root = self.root.resolve()
        raw_root = (root / "data" / "raw").resolve()
        raw_root.mkdir(parents=True, exist_ok=True)

        archive_path = raw_root / safe_name
        if archive_path.exists():
            archive_path = raw_root / f"{archive_path.stem}-{uuid4().hex[:8]}{archive_path.suffix}"
        archive_path.write_bytes(content)

        extracted: list[str] = []
        skipped: list[str] = []
        try:
            with zipfile.ZipFile(archive_path) as archive:
                members: list[tuple[zipfile.ZipInfo, Path]] = []
                for info in archive.infolist():
                    name = info.filename
                    if not name or name.startswith("__MACOSX/") or Path(name).name == ".DS_Store":
                        skipped.append(name)
                        continue
                    target = (raw_root / name).resolve()
                    if target != raw_root and raw_root not in target.parents:
                        raise ValueError(f"Unsafe zip member path: {name}")
                    members.append((info, target))

                for info, target in members:
                    if info.is_dir():
                        target.mkdir(parents=True, exist_ok=True)
                        continue
                    target.parent.mkdir(parents=True, exist_ok=True)
                    with archive.open(info) as source, target.open("wb") as dest:
                        dest.write(source.read())
                    extracted.append(str(target.relative_to(root)))
        except zipfile.BadZipFile as exc:
            archive_path.unlink(missing_ok=True)
            raise ValueError("Uploaded file is not a valid zip archive.") from exc
        except ValueError:
            archive_path.unlink(missing_ok=True)
            raise

        archive_rel = str(archive_path.relative_to(root))
        self.append_timeline({
            "type": "raw_data_zip_uploaded",
            "timestamp": now_iso(),
            "message": archive_rel,
            "payload": {
                "archive": archive_rel,
                "size": len(content),
                "extractedCount": len(extracted),
                "skippedCount": len(skipped),
                "extracted": extracted[:200],
                "truncated": len(extracted) > 200,
            },
        })
        return {
            "archive": archive_rel,
            "size": len(content),
            "extracted": extracted,
            "skipped": skipped,
            "targetDir": "data/raw",
        }

    def ensure_reference_text_derivatives(self) -> None:
        references = self.root / "references"
        if not references.exists():
            return
        for path in references.iterdir():
            if path.is_file() and path.suffix.lower() == ".docx":
                self._write_reference_text_derivative(path)

    def _write_reference_text_derivative(self, path: Path) -> str | None:
        if path.suffix.lower() != ".docx":
            return None
        target = path.with_suffix(path.suffix + ".txt")
        if target.exists() and target.stat().st_mtime >= path.stat().st_mtime:
            return str(target.relative_to(self.root))
        try:
            text = _extract_docx_text(path)
        except Exception:
            return None
        if not text.strip():
            return None
        target.write_text(text, encoding="utf-8")
        return str(target.relative_to(self.root))

    def ensure_builtin_tools(self) -> None:
        script = self.root / "tools" / "read_docx.py"
        if script.exists():
            return
        script.write_text('''from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def extract_docx(path: Path) -> str:
    document = Document(str(path))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        chunks.append(f"\\n[Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\\n", " ") for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    return "\\n\\n".join(chunks)


def main() -> None:
    if len(sys.argv) not in {2, 3}:
        raise SystemExit("Usage: uv run python tools/read_docx.py <path-to-docx> [output.txt]")
    path = Path(sys.argv[1])
    text = extract_docx(path)
    if len(sys.argv) == 3:
        out = Path(sys.argv[2])
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(text + "\\n", encoding="utf-8")
        print(str(out))
    else:
        print(text)


if __name__ == "__main__":
    main()
''', encoding="utf-8")

    def clear_debug_logs(self, scope: str = "main") -> None:
        clear_file(self.main_log_path)
        if scope == "chat":
            self.reset_chat()
            return
        if scope != "all":
            return
        self.reset_workspace()

    def reset_chat(self) -> WorkspaceState:
        existing = self.read_state() or {}
        control_mode = "auto" if existing.get("controlMode") == "auto" else "manual"
        runtime_settings = self.read_runtime_settings()

        for rel in (
            "user",
            "data/processed",
            "plots",
            "tools",
            "runs",
            "reports",
            "training",
            "logs/nodes",
            "state/nodes",
        ):
            _remove_path(self.root / rel)

        for path in (self.state_path, self.main_log_path, self.timeline_path):
            path.unlink(missing_ok=True)

        artifacts_root = self.root / "artifacts"
        if artifacts_root.exists():
            for child in artifacts_root.iterdir():
                if child.name == "knowledge-graph.json":
                    continue
                _remove_path(child)

        self.ensure_layout()
        ts = now_iso()
        state = self._new_workspace_state(control_mode, runtime_settings, ts)
        self.write_state(state)
        self.append_timeline({
            "type": "chat_reset",
            "timestamp": ts,
            "message": "Chat and agent workflow reset",
            "payload": {
                "preserved": [
                    "data/raw",
                    "references",
                    "knowledge_base",
                    "artifacts/knowledge-graph.json",
                    "state/knowledge-graph-build.json",
                    "logs/knowledge-graph-builder.jsonl",
                    "logs/knowledge-reasoning.jsonl",
                ],
                "cleared": [
                    "logs/main.jsonl",
                    "logs/nodes",
                    "state/nodes",
                    "user",
                    "data/processed",
                    "artifacts/* except knowledge-graph.json",
                    "plots",
                    "tools",
                    "runs",
                    "reports",
                    "training",
                ],
            },
        })
        return state

    def reset_workspace(self) -> WorkspaceState:
        existing = self.read_state() or {}
        control_mode = "auto" if existing.get("controlMode") == "auto" else "manual"
        runtime_settings = self.read_runtime_settings()

        for rel in (
            "user",
            "data/raw",
            "data/processed",
            "references",
            "knowledge_base",
            "artifacts",
            "plots",
            "tools",
            "runs",
            "reports",
            "logs",
            "training",
            "state/nodes",
        ):
            _remove_path(self.root / rel)

        for path in (self.state_path, self.knowledge_graph_status_path, self.knowledge_graph_path):
            path.unlink(missing_ok=True)

        self.ensure_layout()
        ts = now_iso()
        state = self._new_workspace_state(control_mode, runtime_settings, ts)
        self.write_state(state)
        self.append_timeline({
            "type": "workspace_reset",
            "timestamp": ts,
            "message": "Workspace reset",
            "payload": {
                "cleared": [
                    "references",
                    "knowledge_base",
                    "logs",
                    "state/nodes",
                    "artifacts",
                    "data",
                    "reports",
                    "runs",
                    "tools",
                    "training",
                ],
            },
        })
        return state

    def _new_workspace_state(self, control_mode: str, runtime_settings: RuntimeSettings, ts: str) -> WorkspaceState:
        return {
            "workspaceId": str(uuid4()),
            "workspacePath": str(self.root),
            "createdAt": ts,
            "updatedAt": ts,
            "mode": control_mode,
            "controlMode": control_mode,
            "pendingControl": None,
            "activeNode": None,
            "activeNodeSessionId": None,
            "completedNodes": [],
            "contractConfirmed": False,
            "finalSummaryConfirmed": False,
            "runtimeSettings": runtime_settings,
        }


def _remove_path(path: Path) -> None:
    if path.is_symlink() or path.is_file():
        path.unlink(missing_ok=True)
    elif path.is_dir():
        shutil.rmtree(path)


def _safe_filename(filename: str) -> str:
    name = Path(filename or "reference").name.strip()
    safe = "".join(ch if ch.isalnum() or ch in {".", "-", "_", " "} else "_" for ch in name)
    safe = safe.strip(" .")
    return safe or f"reference-{uuid4().hex[:8]}"


def _extract_docx_text(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        chunks.append(f"\n[Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    return "\n\n".join(chunks).strip() + "\n"


def _bounded_int(value: Any, *, default: int, minimum: int, maximum: int) -> int:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        parsed = default
    return max(minimum, min(maximum, parsed))


def _sanitize_llm_config(raw: dict[str, Any], *, default_auth_mode: str) -> dict[str, Any]:
    auth = raw.get("authMode") if raw.get("authMode") in {"manual", "sdk-default"} else default_auth_mode
    protocol = raw.get("protocol") if raw.get("protocol") in {"anthropic", "openai-compat"} else None
    context = raw.get("contextWindow") if raw.get("contextWindow") in {"200k", "1m"} else None
    return {
        "authMode": auth,
        "model": raw.get("model") if isinstance(raw.get("model"), str) else "",
        "apiKey": raw.get("apiKey") if isinstance(raw.get("apiKey"), str) else None,
        "baseUrl": raw.get("baseUrl") if isinstance(raw.get("baseUrl"), str) else None,
        "protocol": protocol,
        "contextWindow": context,
    }


def _mask_llm_config_dict(config: dict[str, Any]) -> dict[str, Any]:
    out = dict(config)
    secret = out.get("apiKey")
    if isinstance(secret, str) and secret:
        out["apiKey"] = "****" if len(secret) <= 8 else f"****{secret[-4:]}"
    return out


def _looks_masked_secret(value: str) -> bool:
    return value.startswith("****")
